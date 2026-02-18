from docx import Document
import os

def create_initial_templates():
    """Creates basic placeholder templates for Phase 3."""
    base_path = "src/assets/templates"
    os.makedirs(base_path, exist_ok=True)
    
    # Template 1: Anexo 5.1
    doc = Document()
    doc.add_heading('ANEXO 5.1 - PLAN DE FORMACIÓN', 0)
    doc.add_paragraph('Alumno: {{nombre_alumno}}')
    doc.add_paragraph('Matrícula: {{matricula}}')
    doc.add_paragraph('Proyecto: {{nombre_proyecto}}')
    doc.add_paragraph('Empresa: {{nombre_empresa}}')
    doc.add_paragraph('Mentor Industrial: {{mentor_ue}}')
    doc.add_paragraph('Mentor Académico: {{mentor_ie}}')
    
    doc.add_heading('Materias Inscritas:', level=1)
    # Placeholder for table table
    doc.add_paragraph('{{tabla_materias}}')
    
    doc.save(os.path.join(base_path, "anexo_5_1_template.docx"))
    
    # Template 2: Carta de Asignación
    doc2 = Document()
    doc2.add_heading('CARTA DE ASIGNACIÓN DE MENTOR', 0)
    doc2.add_paragraph('Estimado(a) alumno(a): {{nombre_alumno}}')
    doc2.add_paragraph('Por medio de la presente se le informa que su Mentor Académico asignado para el periodo actual es:')
    doc2.add_paragraph('Profesor: {{mentor_ie}}')
    doc2.add_paragraph('Correo: {{email_mentor_ie}}')
    doc2.add_paragraph('\nAtentamente,\nCoordinación DUAL')
    
    doc2.save(os.path.join(base_path, "carta_asignacion_template.docx"))
    
    print("Templates created successfully.")

if __name__ == "__main__":
    create_initial_templates()
