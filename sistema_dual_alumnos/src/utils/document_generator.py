from docx import Document
import os
import tempfile

def fill_template(template_path, data, output_filename):
    """
    Fills a .docx template with data dictionary values.
    Returns path to generated file.
    """
    try:
        doc = Document(template_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
        
        # Replace in tables (basic)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                             if f"{{{{{key}}}}}" in paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

        tmp_dir = tempfile.gettempdir()
        output_path = os.path.join(tmp_dir, output_filename)
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error generating document: {e}")
        return None

def generate_student_documents(student_data, project_data, mentor_data):
    """Generates Anexo 5.1 and Letter for a student."""
    base_path = "src/assets/templates"
    
    # Prepare Context
    context = {
        "nombre_alumno": f"{student_data.get('nombre')} {student_data.get('ap_paterno')} {student_data.get('ap_materno')}",
        "matricula": student_data.get("matricula"),
        "nombre_proyecto": project_data.get("nombre_proyecto"),
        "nombre_empresa": project_data.get("ue_name", "N/A"), # Needs join in real query
        "mentor_ue": project_data.get("mentor_ue_name", "N/A"),
        "mentor_ie": mentor_data.get("nombre_completo", "Pendiente"),
        "email_mentor_ie": mentor_data.get("email_institucional", "N/A"),
        "tabla_materias": "(Detalle de materias aquí...)" # logic for table rows omitted for brevity
    }
    
    # Generate Anexo
    anexo_path = fill_template(
        os.path.join(base_path, "anexo_5_1_template.docx"),
        context,
        f"Anexo_5.1_{student_data.get('matricula')}.docx"
    )
    
    # Generate Carta
    carta_path = fill_template(
        os.path.join(base_path, "carta_asignacion_template.docx"),
        context,
        f"Carta_Asignacion_{student_data.get('matricula')}.docx"
    )
    
    return [anexo_path, carta_path] if anexo_path and carta_path else []
