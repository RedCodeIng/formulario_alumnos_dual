import os
import io

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def generar_documento_word(template_path_str, output_path_str, context_dict):
    """
    Lee un documento Word de plantilla y reemplaza las etiquetas tipo {{ llave }} 
    por los valores contenidos en context_dict.
    Devuelve (True, "OK") o (False, "Error message").
    """
    if not DOCX_AVAILABLE:
        return False, "La librería python-docx no está instalada."
        
    try:
        doc = Document(template_path_str)
        
        # Helper string replace function
        def replace_text_in_runs(paragraph, mapping):
            # docx text replacement is tricky because a single word can be split across multiple 'runs'
            # A simple approach: join all text, replace, and put it back in the first run, clearing the rest.
            full_text = "".join([run.text for run in paragraph.runs])
            
            modified = False
            for key, val in mapping.items():
                target = f"{{{{ {key} }}}}"
                if target in full_text:
                    full_text = full_text.replace(target, str(val))
                    modified = True
            
            if modified and paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ""
                    
        # 1. Replace in paragraphs
        for paragraph in doc.paragraphs:
            replace_text_in_runs(paragraph, context_dict)
            
        # 2. Replace in tables (Anexos usually have tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_runs(paragraph, context_dict)
                        
        doc.save(output_path_str)
        return True, "Documento generado correctamente."
        
    except Exception as e:
        return False, str(e)
