import os
import copy
import docx
from datetime import datetime
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Pt, Mm
from docxtpl import DocxTemplate, InlineImage

def clone_and_replace(row, text):
    """
    Creates a deep copy of an existing row to ensure we preserve all
    MS Word cell layouts, grid spans, and borders.
    """
    new_tr = copy.deepcopy(row._tr)
    new_row = docx.table._Row(new_tr, row._parent)
    
    # clear all cells
    for cell in new_row.cells:
        for p in cell.paragraphs:
            p.text = ''
            
    # set text in first cell
    if len(new_row.cells) > 0:
        new_row.cells[0].paragraphs[0].text = text
        
    return new_tr

def generate_pdf_from_docx(template_name, context, output_path, template_path=None):
    """
    Generates a PDF from a DOCX template using docxtpl and docx2pdf.
    Dynamically injects {%tr for %} rows to multiply list elements,
    and inserts a page break before the signatures block.
    """
    try:
        if template_path and os.path.exists(template_path):
            t_path = template_path
        else:
            base_dir = os.path.join(os.path.dirname(__file__), '../templates/docs')
            t_path = os.path.join(base_dir, template_name)
            
        if not os.path.exists(t_path):
            return False, f"Plantilla DOCX no encontrada: {t_path}"

        # 2. Preprocess Template to inject dynamic rows and page break
        raw_doc = docx.Document(t_path)
        elaboraron_processed = False
        
        # Only inject these specific tags for Anexo 5.1 Plan de Formacion
        if "Anexo_5.1" in template_name:
            for t_idx, table in enumerate(raw_doc.tables):
                for r_idx, row in enumerate(table.rows):
                    row_text = ''.join(cell.text for cell in row.cells)
                    
                    if '{{ competencia }}' in row_text:
                        row._tr.addprevious(clone_and_replace(row, '{%tr for c in competencias_list %}'))
                        row._tr.addnext(clone_and_replace(row, '{%tr endfor %}'))
                        
                        processed_cells = set()
                        for cell in row.cells:
                            if cell in processed_cells: continue
                            processed_cells.add(cell)
                            for p in cell.paragraphs:
                                if '{{' in p.text:
                                    p.text = p.text.replace('{{ loop_index }}', '{{ loop.index }}') \
                                                   .replace('{{ competencia }}', '{{ c.competencia }}') \
                                                   .replace('{{ asignatura }}', '{{ c.asignatura }}')
                                    
                    elif '{{ actividad }}' in row_text:
                        row._tr.addprevious(clone_and_replace(row, '{%tr for a in actividades_list %}'))
                        row._tr.addnext(clone_and_replace(row, '{%tr endfor %}'))
                        
                        processed_cells = set()
                        for cell in row.cells:
                            if cell in processed_cells: continue
                            processed_cells.add(cell)
                            for p in cell.paragraphs:
                                if '{{' in p.text:
                                    p.text = p.text.replace('{{ loop_index }}', '{{ loop.index }}') \
                                                   .replace('{{ actividad }}', '{{ a.actividad }}') \
                                                   .replace('{{ horas }}', '{{ a.horas }}') \
                                                   .replace('{{ evidencia }}', '{{ a.evidencia }}') \
                                                   .replace('{{ lugar }}', '{{ a.lugar }}') \
                                                   .replace('{{ ponderacion }}', '{{ a.ponderacion }}')
    
                    # Check for ELABORARON (Signatures block) - Insert page break before the signature table
                    if 'ELABORARON' in row_text and not elaboraron_processed:
                        processed_cells = set()
                        for cell in row.cells:
                            if cell in processed_cells: continue
                            processed_cells.add(cell)
                            for p in cell.paragraphs:
                                if 'ELABORARON' in p.text:
                                    p.insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)
                                    elaboraron_processed = True
                                    break
                            if elaboraron_processed: break

        if "Anexo_5.4" in template_name:
            import re
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Pt, Inches, RGBColor

            # Pre-cleanup ghost tags left by user
            for table in raw_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if re.search(r'\{\{\s*(c|act|eval)\.', p.text):
                                p.text = re.sub(r'\{\{\s*(c|act|eval)\.[^}]*\}\}', '', p.text)
            for p in raw_doc.paragraphs:
                if re.search(r'\{\{\s*(c|act|eval)\.', p.text):
                    p.text = re.sub(r'\{\{\s*(c|act|eval)\.[^}]*\}\}', '', p.text)

            # Pre-cleanup ghost template rows (leftover Marco Teorico inside layout table)
            for table in raw_doc.tables:
                rows_to_delete = []
                delete_mode = False
                for row in table.rows:
                    row_text = ''.join(cell.text for cell in row.cells).strip()
                    if '[[TABLA_COMPETENCIAS]]' in row_text:
                        delete_mode = True
                        continue
                    if '[[TABLA_EVALUACION]]' in row_text or '3.- EVALUACIÓN' in row_text or 'EVALUACION' in row_text:
                        delete_mode = False
                        continue
                    if delete_mode:
                        rows_to_delete.append(row)
                        
                for r in rows_to_delete:
                    # Remove the row element from its parent gracefully
                    r._element.getparent().remove(r._element)

            def set_cell_background(cell, fill="BFBFBF"):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                tcPr.append(shd)

            def make_header_run(cell, text, font_size=10):
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Helvetica'
                run.font.size = Pt(font_size)
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                set_cell_background(cell, "BFBFBF")
                
            def make_data_run(cell, text):
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = 'Helvetica'
                run.font.size = Pt(10)

            competencias_added = False
            evaluaciones_added = False

            for layout_table in raw_doc.tables:
                for row in layout_table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if '[[TABLA_COMPETENCIAS]]' in p.text:
                                p.text = p.text.replace('[[TABLA_COMPETENCIAS]]', '')
                                if not competencias_added:
                                    competencias_added = True
                                    # Create Tabla Competencias natively inside the cell
                                    table = cell.add_table(rows=1, cols=3)
                                    table.style = 'Table Grid'
                                    
                                    cells = table.rows[0].cells
                                    make_header_run(cells[0], "No.")
                                    make_header_run(cells[1], "COMPETENCIAS A DESARROLLAR")
                                    make_header_run(cells[2], "ASIGNATURAS")
                                    
                                    # Set widths (approximate percentages to fit 7.5 inch narrow margins)
                                    widths_comp = [0.5, 3.5, 3.5]
                                    for table_row in table.rows:
                                        for idx, width in enumerate(widths_comp):
                                            table_row.cells[idx].width = Inches(width)
                                    
                                    # Add data rows
                                    mt_parts = []
                                    desc_parts = []
                                    for comp in context.get("lista_competencias", []):
                                        # Main Row (No, Competencia, Asignaturas)
                                        row_comp = table.add_row()
                                        for idx, width in enumerate(widths_comp):
                                            row_comp.cells[idx].width = Inches(width)
                                        make_data_run(row_comp.cells[0], str(comp.get("numero_consecutivo", "")))
                                        make_data_run(row_comp.cells[1], comp.get("competencia_desarrollada", ""))
                                        make_data_run(row_comp.cells[2], comp.get("asignaturas_cubre", ""))
                                        row_comp.cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                        
                                        if comp.get("conocimientos_teoricos"):
                                            mt_parts.append(f"{comp.get('competencia_desarrollada')}:\n{comp.get('conocimientos_teoricos')}")
                                        if comp.get("descripcion_actividades"):
                                            desc_parts.append(f"{comp.get('competencia_desarrollada')}:\n{comp.get('descripcion_actividades')}")
                                            
                                    p1 = cell.add_paragraph("\xa0") # Forced space to prevent table fusion
                                    p1.paragraph_format.space_after = Pt(12)
                                    p1_run = p1.runs[0]
                                    p1_run.font.size = Pt(2)
                                    
                                    # Table 2: Marco Teorico
                                    table_mt = cell.add_table(rows=2, cols=1)
                                    table_mt.style = 'Table Grid'
                                    mt_header = table_mt.rows[0].cells[0]
                                    make_header_run(mt_header, "MARCO TEÓRICO O ANTECEDENTES")
                                    mt_header.width = Inches(7.5)
                                    make_data_run(table_mt.rows[1].cells[0], "\n\n".join(mt_parts) if mt_parts else "")
                                    table_mt.rows[1].cells[0].width = Inches(7.5)
                                    
                                    p2 = cell.add_paragraph("\xa0") # Forced space to prevent table fusion
                                    p2.paragraph_format.space_after = Pt(12)
                                    p2_run = p2.runs[0]
                                    p2_run.font.size = Pt(2)
                                    
                                    # Table 3: Descripcion
                                    table_desc = cell.add_table(rows=2, cols=1)
                                    table_desc.style = 'Table Grid'
                                    desc_header = table_desc.rows[0].cells[0]
                                    make_header_run(desc_header, "DESCRIPCIÓN DE LAS ACTIVIDADES REALIZADAS")
                                    desc_header.width = Inches(7.5)
                                    make_data_run(table_desc.rows[1].cells[0], "\n\n".join(desc_parts) if desc_parts else "")
                                    table_desc.rows[1].cells[0].width = Inches(7.5)

                            if '[[TABLA_EVALUACION]]' in p.text:
                                p.text = p.text.replace('[[TABLA_EVALUACION]]', '')
                                if not evaluaciones_added:
                                    evaluaciones_added = True
                                    evals = context.get("evaluaciones", [])
                                    for ev in evals:
                                        # Header Row 1 (COMPETENCIA)
                                        table = cell.add_table(rows=2, cols=9)
                                        table.style = 'Table Grid'
                                        
                                        # Row 0: COMPETENCIA title
                                        row0 = table.rows[0]
                                        make_header_run(row0.cells[0], "C")
                                        run_comp = row0.cells[0].paragraphs[0].add_run("OMPETENCIA") # Visual trick
                                        run_comp.bold = True
                                        row0.cells[0].merge(row0.cells[1])
                                        
                                        p_target = row0.cells[2].paragraphs[0]
                                        run_targ = p_target.add_run(ev.get("competencia_alcanzada", ""))
                                        run_targ.font.name = 'Helvetica'
                                        run_targ.font.size = Pt(10)
                                        for i in range(2, 9):
                                            row0.cells[2].merge(row0.cells[i])
                                            
                                        # Row 1 and 2 logic: Merge first, then add headers
                                        row1 = table.rows[1]
                                        row2 = table.add_row()
                                        
                                        # Set absolute column widths for all 9 columns to fit 7.5 margins
                                        widths_eval = [1.5, 1.5, 0.7, 0.4, 0.4, 0.4, 0.4, 0.4, 1.8]
                                        for table_row in table.rows:
                                            for idx, width in enumerate(widths_eval):
                                                table_row.cells[idx].width = Inches(width)

                                        # Vertical merges
                                        c_act = row1.cells[0].merge(row2.cells[0])
                                        c_evi = row1.cells[1].merge(row2.cells[1])
                                        c_hor = row1.cells[2].merge(row2.cells[2])
                                        c_fir = row1.cells[8].merge(row2.cells[8])
                                        
                                        # Horizontal merge for NIVEL DE DESEMPEÑO
                                        c_niv = row1.cells[3].merge(row1.cells[7])
                                        
                                        # Now write headers to the safely merged cell containers
                                        make_header_run(c_act, "ACTIVIDADES")
                                        make_header_run(c_evi, "EVIDENCIAS O PRODUCTOS")
                                        make_header_run(c_hor, "HORAS DE DEDICACIÓN")
                                        make_header_run(c_niv, "NIVEL DE DESEMPEÑO")
                                        make_header_run(c_fir, "NOMBRE,\nFIRMA Y\nFECHA DE\nEVALUACIÓN\nDEL\nMENTOR DE\nLA UE", font_size=7.5)
                                        
                                        # Row 2 headers
                                        make_header_run(row2.cells[3], "0%")
                                        make_header_run(row2.cells[4], "70%")
                                        make_header_run(row2.cells[5], "80%")
                                        make_header_run(row2.cells[6], "90%")
                                        make_header_run(row2.cells[7], "100%")
                                        
                                        for act_idx, act in enumerate(ev.get("actividades", [])):
                                            act_row = table.add_row()
                                            for idx, width in enumerate(widths_eval):
                                                act_row.cells[idx].width = Inches(width)
                                                
                                            make_data_run(act_row.cells[0], act.get("descripcion_actividad", ""))
                                            make_data_run(act_row.cells[1], act.get("evidencia", ""))
                                            make_data_run(act_row.cells[2], str(act.get("horas", "")))
                                            
                                            make_data_run(act_row.cells[3], act.get("p0", ""))
                                            make_data_run(act_row.cells[4], act.get("p70", ""))
                                            make_data_run(act_row.cells[5], act.get("p80", ""))
                                            make_data_run(act_row.cells[6], act.get("p90", ""))
                                            make_data_run(act_row.cells[7], act.get("p100", ""))
                                            
                                            # Center align all levels
                                            for i in range(2, 8):
                                                act_row.cells[i].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            if act_idx == 0:
                                                make_data_run(act_row.cells[8], ev.get("firma_y_fecha", ""))
                                                act_row.cells[8].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                                # Fix sizing for signature block if needed, but standard 10 is fine
                                            else:
                                                # Merge cells downward for the signature column
                                                table.cell(2, 8).merge(act_row.cells[8])

                                        # Add spacing paragraph between evaluations
                                        cell.add_paragraph("")

        injected_path = output_path.replace('.pdf', '_INJECTED.docx')
        if not injected_path.endswith('.docx'): injected_path += '_INJECTED.docx'
        raw_doc.save(injected_path)

        # 3. Render Template using docxtpl on the injected file
        doc = DocxTemplate(injected_path)
        
        # Helper to recursively find strings starting with IMAGE_PATH: and convert to InlineImage
        def _process_images_in_context(ctx_item):
            if isinstance(ctx_item, dict):
                for k, v in ctx_item.items():
                    if isinstance(v, str) and v.startswith("IMAGE_PATH:"):
                        img_path = v.split("IMAGE_PATH:")[1]
                        if os.path.exists(img_path):
                            ctx_item[k] = InlineImage(doc, img_path, width=Mm(40)) # Typical 4cm donut
                        else:
                            ctx_item[k] = "" # fallback if missing
                    else:
                        _process_images_in_context(v)
            elif isinstance(ctx_item, list):
                for item in ctx_item:
                    _process_images_in_context(item)

        _process_images_in_context(context)
        
        # docxtpl uses jinja2 under the hood, we just pass the context
        doc.init_docx()
        with open('final_jinja_src.txt', 'w', encoding='utf-8') as f:
            f.write(doc.patch_xml(doc.get_xml()))
        doc.render(context)
            
        # 4. Save filled DOCX temporarily
        temp_docx_path = output_path.replace('.pdf', '.docx')
        if not temp_docx_path.endswith('.docx'):
             temp_docx_path += '.docx'
             
        doc.save(temp_docx_path)
        
        # 6. Convert to PDF using docx2pdf (Requires MS Word installed)
        try:
            from docx2pdf import convert
            convert(temp_docx_path, output_path)
            
            # Clean up temporary DOCX files to save resources
            try:
                os.remove(injected_path)
                os.remove(temp_docx_path)
            except Exception as e:
                print(f"Cleanup warning: {e}")
            
            return True, f"Documento generado exitosamente en: {output_path}"
            
        except ImportError:
            # If docx2pdf fails or isn't available, we return the filled DOCX
            temp_pdf_path = output_path
            return True, f"DOCX rellenado exitosamente en: {temp_docx_path} (No se pudo convertir a PDF automáticamente)"
        except Exception as e:
            return True, f"DOCX guardado en {temp_docx_path}. Error al convertir a PDF: {e}"

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise e
